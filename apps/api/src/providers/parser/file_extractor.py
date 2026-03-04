"""
File text extraction utilities.

Converts various binary file formats to plain text for indexing.
Images are processed via Claude Vision API.
"""

import csv
import io
import json
from pathlib import Path

# extension → (parser_type_for_chunking, category)
SUPPORTED_EXTENSIONS: dict[str, str] = {
    # Plain text
    ".txt": "text",
    ".md": "markdown",
    ".markdown": "markdown",
    # Documents
    ".pdf": "text",
    ".docx": "text",
    # Spreadsheets
    ".csv": "text",
    ".xlsx": "text",
    ".xls": "text",
    # Jupyter Notebook
    ".ipynb": "markdown",
    # Images (extracted via Claude Vision)
    ".jpg": "text",
    ".jpeg": "text",
    ".png": "text",
    ".gif": "text",
    ".webp": "text",
    ".bmp": "text",
    ".tiff": "text",
    ".tif": "text",
}

# Anthropic-native image MIME types
_ANTHROPIC_MEDIA_TYPES: dict[str, str] = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".gif": "image/gif",
    ".webp": "image/webp",
}
# Formats that need Pillow conversion → PNG before sending to Claude
_NEEDS_CONVERSION = {".bmp", ".tiff", ".tif"}

IMAGE_EXTENSIONS = set(_ANTHROPIC_MEDIA_TYPES) | _NEEDS_CONVERSION


def get_file_type(filename: str) -> str | None:
    """Return the chunker type for the file, or None if unsupported."""
    ext = Path(filename).suffix.lower()
    return SUPPORTED_EXTENSIONS.get(ext)


async def extract_text(
    filename: str,
    data: bytes,
    anthropic_api_key: str = "",
    high_quality: bool = False,
) -> tuple[str, str]:
    """
    Extract plain text from a file's raw bytes.

    Returns:
        (text_content, parser_type)  — parser_type fed to create_parser()
    Raises:
        ValueError: unsupported extension
        RuntimeError: missing optional dependency
    """
    ext = Path(filename).suffix.lower()
    parser_type = SUPPORTED_EXTENSIONS.get(ext)
    if parser_type is None:
        raise ValueError(f"Unsupported file type: {ext!r}")

    if ext in (".txt", ".md", ".markdown"):
        text = data.decode("utf-8", errors="replace")
    elif ext == ".pdf":
        if high_quality:
            if not anthropic_api_key:
                raise ValueError("高质量模式需要配置 ANTHROPIC_API_KEY")
            text = await _pdf_vision(data, anthropic_api_key)
        else:
            text = _pdf(data)
    elif ext == ".docx":
        text = _docx(data)
    elif ext == ".csv":
        text = _csv(data)
    elif ext in (".xlsx", ".xls"):
        text = _excel(data, ext)
    elif ext == ".ipynb":
        text = _notebook(data)
    elif ext in IMAGE_EXTENSIONS:
        text = await _image(data, ext, anthropic_api_key)
    else:
        raise ValueError(f"Unhandled extension: {ext!r}")

    # PostgreSQL UTF-8 rejects null bytes — strip them from all sources
    return text.replace('\x00', ''), parser_type


# ---------------------------------------------------------------------------
# Format extractors
# ---------------------------------------------------------------------------

def _pdf(data: bytes) -> str:
    try:
        import pypdf
    except ImportError:
        raise RuntimeError("PDF support requires: pip install pypdf")
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            raise RuntimeError("PDF 已加密，无法提取文本。请解密后重新上传，或使用高质量模式。")
        parts = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(p for p in parts if p.strip())
    except RuntimeError:
        raise
    except Exception as e:
        raise RuntimeError(f"PDF 解析失败: {e}") from e


async def _pdf_vision(data: bytes, api_key: str, max_pages: int = 50) -> str:
    """Render each PDF page to PNG and extract content via Claude Vision."""
    import base64
    import anthropic

    try:
        import fitz  # pymupdf
    except ImportError:
        raise RuntimeError("高质量 PDF 模式需要: pip install pymupdf")

    doc = fitz.open(stream=data, filetype="pdf")
    total_pages = len(doc)
    pages_to_process = min(total_pages, max_pages)

    client = anthropic.AsyncAnthropic(api_key=api_key)
    parts: list[str] = []

    for i in range(pages_to_process):
        page = doc[i]
        # 2x zoom for sharper text recognition
        mat = fitz.Matrix(2, 2)
        pix = page.get_pixmap(matrix=mat)
        img_data = pix.tobytes("png")

        response = await client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=2000,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/png",
                            "data": base64.standard_b64encode(img_data).decode(),
                        },
                    },
                    {
                        "type": "text",
                        "text": (
                            f"这是 PDF 第 {i + 1}/{total_pages} 页。"
                            "请提取页面中的所有内容：正文、标题、数学公式（用 LaTeX 格式表示）、"
                            "代码块（保留缩进）、表格（用 Markdown 格式）。"
                            "忽略页眉、页脚、页码。直接输出内容，不要添加解释。"
                        ),
                    },
                ],
            }],
        )
        parts.append(response.content[0].text)

    if total_pages > max_pages:
        parts.append(f"\n[文档共 {total_pages} 页，高质量模式已处理前 {max_pages} 页]")

    return "\n\n---\n\n".join(parts)


def _docx(data: bytes) -> str:
    try:
        import docx
    except ImportError:
        raise RuntimeError("Word support requires: pip install python-docx")
    doc = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    return "\n".join(" | ".join(row) for row in reader if any(row))


def _excel(data: bytes, ext: str) -> str:
    if ext == ".xls":
        try:
            import xlrd
        except ImportError:
            raise RuntimeError("Legacy Excel (.xls) support requires: pip install xlrd")
        wb = xlrd.open_workbook(file_contents=data)
        parts: list[str] = []
        for sheet in wb.sheets():
            parts.append(f"## Sheet: {sheet.name}")
            for r in range(sheet.nrows):
                row = [str(sheet.cell_value(r, c)) for c in range(sheet.ncols)]
                if any(v.strip() for v in row):
                    parts.append(" | ".join(row))
        return "\n\n".join(parts)
    else:
        try:
            import openpyxl
        except ImportError:
            raise RuntimeError("Excel support requires: pip install openpyxl")
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"## Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(v) for v in row if v is not None]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)


def _notebook(data: bytes) -> str:
    nb = json.loads(data.decode("utf-8", errors="replace"))
    parts: list[str] = []
    for cell in nb.get("cells", []):
        source = "".join(cell.get("source", []))
        if not source.strip():
            continue
        if cell.get("cell_type") == "markdown":
            parts.append(source)
        elif cell.get("cell_type") == "code":
            parts.append(f"```python\n{source}\n```")
            for output in cell.get("outputs", []):
                if output.get("output_type") in ("stream", "display_data", "execute_result"):
                    out_text = output.get("text", [])
                    if isinstance(out_text, list):
                        out_text = "".join(out_text)
                    if out_text.strip():
                        parts.append(f"Output:\n{out_text.strip()}")
    return "\n\n".join(parts)


async def _image(data: bytes, ext: str, api_key: str) -> str:
    """Extract text/description from image via Claude Vision (claude-haiku)."""
    import base64
    import anthropic

    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY required for image extraction")

    # Convert unsupported formats to PNG via Pillow
    if ext in _NEEDS_CONVERSION:
        try:
            from PIL import Image
        except ImportError:
            raise RuntimeError("BMP/TIFF support requires: pip install Pillow")
        img = Image.open(io.BytesIO(data))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        data = buf.getvalue()
        media_type = "image/png"
    else:
        media_type = _ANTHROPIC_MEDIA_TYPES[ext]

    client = anthropic.AsyncAnthropic(api_key=api_key)
    response = await client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=2000,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": base64.standard_b64encode(data).decode(),
                    },
                },
                {
                    "type": "text",
                    "text": (
                        "请提取这张图片中的所有文字内容。"
                        "如果图片中没有文字，请详细描述图片的主要内容，"
                        "包括图表、公式、图形等信息。"
                    ),
                },
            ],
        }],
    )
    return response.content[0].text
